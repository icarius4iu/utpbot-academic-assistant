"""
Servicio de Google Gemini para el sistema UTPBot.
Usa el SDK moderno 'google-genai' (v1.x) que llama a la API v1,
compatible con gemini-2.0-flash, gemini-1.5-flash, etc. en tier gratuito.
Soporta análisis de documentos: PDF (nativo), DOCX, XLSX, TXT, CSV.
"""

import os
import io
import json
import re
import time
import base64
import warnings
from dotenv import load_dotenv

# Silenciar los molestos warnings de Pydantic al serializar objetos Content en el SDK de Gemini
warnings.filterwarnings("ignore", category=UserWarning, module="pydantic")
from typing import List, Dict, Tuple, Optional
from google import genai
from google.genai import types

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

# Modelos disponibles en la API v1. Prioridad a modelos PRO para mayor calidad
MODELOS_FALLBACK = [
    os.getenv("GEMINI_MODEL", "gemini-2.5-pro"),  # 1. Pro más reciente (Alta calidad)
    "gemini-1.5-pro-latest",                      # 2. Pro estable
    "gemini-2.5-flash",                           # 3. Flash ultrarrápido
    "gemini-2.0-flash-001",
    "gemini-flash-latest",
]

# Cliente global (se inicializa una vez)
_client = None

def _get_client() -> genai.Client:
    global _client
    if _client is None:
        if not GEMINI_API_KEY:
            raise ValueError(
                "GEMINI_API_KEY no configurada. "
                "Ve a https://aistudio.google.com/app/apikey para obtener tu key."
            )
        _client = genai.Client(api_key=GEMINI_API_KEY)
        print("[OK] Cliente Gemini (google-genai SDK) inicializado.")
    return _client


def _extraer_texto_docx(file_bytes: bytes) -> str:
    """Extrae texto de un archivo Word (.docx) usando python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        # Incluir también tablas
        for table in doc.tables:
            for row in table.rows:
                fila = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if fila:
                    parrafos.append(fila)
        return "\n".join(parrafos)
    except Exception as e:
        print(f"[WARN] Error extrayendo DOCX: {e}")
        return "(No se pudo extraer el contenido del documento Word)"


def _extraer_texto_xlsx(file_bytes: bytes) -> str:
    """Extrae datos de un archivo Excel (.xlsx) usando openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        resultado = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            resultado.append(f"### Hoja: {sheet_name}")
            filas_con_datos = 0
            for row in ws.iter_rows(values_only=True):
                valores = [str(v) if v is not None else "" for v in row]
                if any(v.strip() for v in valores):
                    resultado.append(" | ".join(valores))
                    filas_con_datos += 1
                if filas_con_datos > 200:  # Límite para no exceder tokens
                    resultado.append("... (más filas omitidas por longitud)")
                    break
        wb.close()
        return "\n".join(resultado)
    except Exception as e:
        print(f"[WARN] Error extrayendo XLSX: {e}")
        return "(No se pudo extraer el contenido del archivo Excel)"


def _extraer_texto_pdf(file_bytes: bytes) -> str:
    """Extrae texto de PDF usando PyMuPDF como fallback si Gemini no lo soporta inline."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        partes = []
        for page in doc:
            t = page.get_text("text")
            if t.strip():
                partes.append(t)
        doc.close()
        return "\n".join(partes)
    except Exception as e:
        print(f"[WARN] Error extrayendo PDF con PyMuPDF: {e}")
        return ""

def agendar_tiempo_estudio(
    titulo: str,
    fecha_inicio: str,
    fecha_fin: str,
    descripcion: str = ""
) -> str:
    """
    Agenda una sesión de estudio en el Google Calendar y envía una confirmación por Telegram.
    
    Args:
        titulo: Título descriptivo de la materia o tema a estudiar.
        fecha_inicio: Fecha y hora de inicio en formato ISO 8601 (ej. "2026-06-03T16:00:00").
        fecha_fin: Fecha y hora de fin en formato ISO 8601 (ej. "2026-06-03T18:00:00").
        descripcion: Detalles o temas a estudiar durante la sesión.
    """
    return "agendar_tiempo_estudio_call"


class GeminiService:
    """
    Servicio para interactuar con la API de Google Gemini (SDK google-genai v1.x).
    Incluye fallback automático entre modelos y retry en cuota agotada.
    Soporta análisis de documentos: PDF (nativo Gemini), DOCX, XLSX, TXT, CSV.
    """

    async def generar_respuesta(
        self,
        system_prompt: str,
        mensaje_usuario: str,
        historial: List[Dict] = None,
        datos_usuario: dict = None,
        file_name: Optional[str] = None,
        file_mime: Optional[str] = None,
        file_data: Optional[str] = None
    ) -> Tuple[str, List[str]]:
        """
        Genera una respuesta del chatbot con fallback entre modelos y soporte de Function Calling.
        Soporta documentos adjuntos: PDF (nativo), DOCX, XLSX, TXT, CSV.
        """
        client = _get_client()
        messages = self._construir_messages(
            mensaje_usuario, historial, file_name, file_mime, file_data
        )
        modelos = list(dict.fromkeys(MODELOS_FALLBACK))
        ultimo_error = None

        for nombre in modelos:
            try:
                print(f"🤖 Intentando modelo: {nombre}")

                response = client.models.generate_content(
                    model=nombre,
                    contents=messages,
                    config=types.GenerateContentConfig(
                        system_instruction=system_prompt,
                        temperature=0.7,
                        max_output_tokens=8192,
                        tools=[agendar_tiempo_estudio],
                        automatic_function_calling=types.AutomaticFunctionCallingConfig(
                            disable=True
                        )
                    )
                )

                # Interceptar invocación de herramienta/función (Function Calling)
                # El SDK google-genai puede exponer las llamadas como response.function_calls
                # o dentro de response.candidates[0].content.parts
                function_call = None

                # Método 1: propiedad directa del SDK
                try:
                    fc_list = response.function_calls
                    if fc_list:
                        function_call = fc_list[0]
                except Exception:
                    pass

                # Método 2: inspeccionar parts de los candidates
                if function_call is None:
                    try:
                        for cand in response.candidates or []:
                            for part in cand.content.parts or []:
                                if hasattr(part, "function_call") and part.function_call:
                                    function_call = part.function_call
                                    break
                            if function_call:
                                break
                    except Exception:
                        pass

                if function_call and function_call.name == "agendar_tiempo_estudio":
                    args = dict(function_call.args)
                    t_titulo = args.get("titulo")
                    t_inicio = args.get("fecha_inicio")
                    t_fin = args.get("fecha_fin")
                    t_desc = args.get("descripcion", "")

                    print(f"📅 [TOOL] Ejecutando agendar_tiempo_estudio para: {t_titulo} ({t_inicio} - {t_fin})")

                    # Importación perezosa para evitar dependencias circulares
                    from services.calendar_service import calendar_service
                    from services.telegram_service import telegram_service

                    try:
                        # 1. Agendar en Google Calendar
                        await calendar_service.crear_evento_estudio(
                            titulo=t_titulo,
                            fecha_inicio=t_inicio,
                            fecha_fin=t_fin,
                            descripcion=t_desc
                        )

                        # 2. Enviar confirmación a Telegram
                        await telegram_service.enviar_confirmacion_estudio(
                            titulo=t_titulo,
                            fecha_inicio=t_inicio,
                            fecha_fin=t_fin,
                            descripcion=t_desc
                        )

                        # Formatear la fecha para que sea legible en la respuesta de chat
                        try:
                            fecha_str = t_inicio.split("T")[0]
                            h_inicio = t_inicio.split("T")[1][:5]
                            h_fin = t_fin.split("T")[1][:5]
                            y, m, d = fecha_str.split("-")
                            fecha_legible = f"{d}/{m}/{y}"
                        except Exception:
                            fecha_legible = t_inicio
                            h_inicio = "Ver"
                            h_fin = "Ver"

                        exito_msg = (
                            f"Sesion de estudio agendada con exito.\n\n"
                            f"He reservado el bloque en tu Google Calendar y te he enviado una confirmacion instantanea a tu Telegram:\n\n"
                            f"Materia: {t_titulo}\n"
                            f"Fecha: {fecha_legible}\n"
                            f"Horario: {h_inicio} - {h_fin}\n\n"
                            f"Detalles: {t_desc or 'Sin descripcion adicional.'}\n\n"
                            f"Mucho exito en tu preparacion academica."
                        )
                        return exito_msg, ["Cual es mi horario?", "Cuales son mis notas?", "Cuando es mi proximo examen?"]
                    except Exception as tool_err:
                        print(f"[ERROR TOOL] Falló ejecución de herramienta: {tool_err}")
                        return (
                            f"Error al agendar la sesion de estudio: {tool_err}\n\n"
                            "Por favor, verifica que tu calendario este compartido correctamente con la cuenta de servicio o intenta nuevamente.",
                            ["Cual es mi horario?", "Cuales son mis notas?", "Cuando es mi proximo examen?"]
                        )

                texto = response.text.strip()

                respuesta_limpia = self._limpiar_respuesta(texto)
                sugerencias = self._extraer_sugerencias(texto)
                print(f"[OK] Respuesta con: {nombre}")
                return respuesta_limpia, sugerencias

            except Exception as e:
                err = str(e)
                ultimo_error = err
                if "429" in err or "quota" in err.lower() or "rate" in err.lower():
                    print(f"[WARN] Cuota agotada en '{nombre}', probando siguiente...")
                    time.sleep(1)
                    continue
                elif "404" in err or "not found" in err.lower():
                    print(f"[WARN] Modelo '{nombre}' no disponible, probando siguiente...")
                    continue
                elif "503" in err or "unavailable" in err.lower() or "demand" in err.lower() or "500" in err:
                    print(f"[WARN] Modelo '{nombre}' sobrecargado (503), probando siguiente...")
                    continue
                else:
                    print(f"[ERROR] Error en '{nombre}': {e}")
                    continue # Siempre intentar otro modelo si hay error

        print(f"[ERROR] Todos los modelos fallaron. Último: {ultimo_error}")

        if "429" in str(ultimo_error) or "quota" in str(ultimo_error).lower():
            msg = (
                "⚠️ **Error 429: Cuota o Facturación.** El servicio de Gemini no procesó la "
                "solicitud porque el proyecto de la API Key no tiene saldo o llegó al límite. "
                f"Detalle: {ultimo_error[:150]}"
            )
        else:
            msg = f"❌ Error llamando a Gemini: {ultimo_error[:150]}"

        return msg, ["¿Cuál es mi horario?", "¿Cuáles son mis notas?", "¿Cuándo es mi próximo examen?"]

    def _construir_messages(
        self,
        mensaje_usuario: str,
        historial: List[Dict],
        file_name: Optional[str] = None,
        file_mime: Optional[str] = None,
        file_data: Optional[str] = None
    ) -> list:
        """
        Construye la lista de mensajes para el SDK google-genai.
        Maneja documentos adjuntos según su tipo MIME:
        - PDF: envío nativo a Gemini como bytes inline
        - DOCX/XLSX/TXT/CSV: extracción de texto local + inyección en mensaje
        """
        messages = []

        # Historial previo
        if historial:
            for msg in historial:
                role = "user" if msg.get("rol") == "user" else "model"
                messages.append(
                    types.Content(
                        role=role,
                        parts=[types.Part.from_text(text=msg.get("contenido", ""))]
                    )
                )

        # Mensaje actual del usuario
        lista_partes = []
        texto_doc_extraido = ""

        if file_data and file_mime:
            try:
                file_bytes = base64.b64decode(file_data)
                mime_lower = file_mime.lower()

                if "pdf" in mime_lower:
                    # PDF: Gemini lo procesa nativamente como bytes inline
                    print(f"📄 Procesando PDF nativo ({len(file_bytes)} bytes)")
                    lista_partes.append(
                        types.Part.from_bytes(
                            data=file_bytes,
                            mime_type="application/pdf"
                        )
                    )

                elif "wordprocessingml" in mime_lower or "msword" in mime_lower or (file_name and file_name.lower().endswith(".docx")):
                    # Word (.docx): extraer texto con python-docx
                    print(f"📝 Procesando DOCX: {file_name}")
                    texto_doc_extraido = _extraer_texto_docx(file_bytes)

                elif "spreadsheetml" in mime_lower or "excel" in mime_lower or (file_name and file_name.lower().endswith(".xlsx")):
                    # Excel (.xlsx): extraer datos con openpyxl
                    print(f"📊 Procesando XLSX: {file_name}")
                    texto_doc_extraido = _extraer_texto_xlsx(file_bytes)

                elif "csv" in mime_lower or (file_name and file_name.lower().endswith(".csv")):
                    # CSV: decodificar texto
                    print(f"📋 Procesando CSV: {file_name}")
                    texto_doc_extraido = file_bytes.decode("utf-8", errors="replace")[:8000]

                elif "text" in mime_lower or (file_name and file_name.lower().endswith(".txt")):
                    # TXT: decodificar texto
                    print(f"📃 Procesando TXT: {file_name}")
                    texto_doc_extraido = file_bytes.decode("utf-8", errors="replace")[:8000]

                else:
                    # Tipo desconocido: intentar como bytes inline
                    print(f"⚠️ Tipo MIME no reconocido ({file_mime}), intentando como bytes inline")
                    lista_partes.append(
                        types.Part.from_bytes(
                            data=file_bytes,
                            mime_type=file_mime
                        )
                    )

            except Exception as e:
                print(f"[ERROR] Error procesando archivo adjunto: {e}")
                texto_doc_extraido = f"(Error al procesar el archivo: {e})"

        # Construir el texto del mensaje final
        texto_mensaje_final = mensaje_usuario
        if texto_doc_extraido:
            nombre_archivo = file_name or "documento adjunto"
            texto_mensaje_final = (
                f"[DOCUMENTO ADJUNTO: {nombre_archivo}]\n"
                f"{'='*50}\n"
                f"{texto_doc_extraido[:6000]}\n"
                f"{'='*50}\n\n"
                f"Consulta del usuario sobre este documento:\n{mensaje_usuario}"
            )

        lista_partes.append(types.Part.from_text(text=texto_mensaje_final))
        messages.append(types.Content(role="user", parts=lista_partes))
        return messages

    def _extraer_sugerencias(self, respuesta: str) -> List[str]:
        """Extrae el JSON de sugerencias del final de la respuesta."""
        default = ["¿Cuál es mi horario?", "¿Cuáles son mis notas?", "¿Cuándo es mi próximo examen?"]
        try:
            patron = r'\{[\s]*"sugerencias"[\s]*:[\s]*\[.*?\][\s]*\}'
            matches = re.findall(patron, respuesta, re.DOTALL)
            if matches:
                data = json.loads(matches[-1])
                sugs = data.get("sugerencias", default)
                return (sugs + default)[:3]
        except Exception:
            pass
        return default

    def _limpiar_respuesta(self, respuesta: str) -> str:
        """Quita el bloque JSON de sugerencias del texto visible."""
        patron = r'\{[\s]*"sugerencias"[\s]*:[\s]*\[.*?\][\s]*\}'
        return re.sub(patron, '', respuesta, flags=re.DOTALL).strip()

    def categorizar_pregunta(self, pregunta: str) -> str:
        """Categoriza la pregunta usando palabras clave para el FAQ_Log."""
        p = pregunta.lower()
        cats = {
            "horarios":   ["horario", "hora", "clase", "aula"],
            "notas":      ["nota", "calificación", "promedio", "aprobé"],
            "examenes":   ["examen", "parcial", "final", "sustitutorio"],
            "asistencia": ["asistencia", "falta", "tardanza", "ausencia"],
            "trabajos":   ["trabajo", "tarea", "entregable", "pendiente"],
            "proyectos":  ["proyecto", "sustentación", "grupo"],
            "calendario": ["evento", "calendario", "fecha", "feriado"],
            "cursos":     ["curso", "materia", "crédito", "matrícula"],
            "docente":    ["profesor", "docente", "sección", "alumnos"],
            "documentos": ["documento", "pdf", "archivo", "excel", "word", "adjunto"],
            "utp_info":   ["utp", "universidad", "campus", "sede", "requisito", "admisión"],
        }
        for cat, palabras in cats.items():
            if any(w in p for w in palabras):
                return cat
        return "general"


# Instancia global
gemini_service = GeminiService()
